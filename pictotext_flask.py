#!/usr/bin/env python3
"""
PicToText - Полная версия с обработкой всех типов файлов
"""

import os
import sys
import tempfile
import logging
import pythoncom
import win32com.client
import time
from threading import Lock
from datetime import datetime
from flask import Flask, request, jsonify, redirect, url_for, flash, render_template, session
from werkzeug.utils import secure_filename
import cv2
import pytesseract
import numpy as np
from pymongo import MongoClient
from bson.objectid import ObjectId
import pdfplumber
from docx import Document
from pdf2image import convert_from_path
from PIL import Image, UnidentifiedImageError
import io
import re
from statistics import mean
import zipfile
import subprocess
from tempfile import mkdtemp, NamedTemporaryFile

# Настройка путей
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
POPPLER_PATH = r'C:\Program Files\poppler-24.08.0\Library\bin'

# Настройка временной директории
if sys.platform == 'win32':
    tempfile.tempdir = os.path.join(os.environ['USERPROFILE'], 'temp_ocrtemp')
    os.makedirs(tempfile.tempdir, exist_ok=True)

app = Flask(__name__)
app.secret_key = 'your-secret-key-here'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10 MB

# Конфигурация
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt', 'jpg', 'jpeg', 'png', 'tiff', 'bmp', 'webp'}

# Настройка MongoDB
MONGO_URL = 'mongodb://localhost:27017/'
DB_NAME = 'ocr_database'
client = MongoClient(MONGO_URL)
db = client[DB_NAME]

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Проверка наличия win32com
try:
    import win32com.client

    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False
    logger.warning("pywin32 not installed, DOC file processing will be limited")

# Глобальная блокировка для COM
com_lock = Lock()


class AccuracyCalculator:
    @staticmethod
    def calculate_from_confidence(conf_data):
        valid_confs = [c for c in conf_data if c > 0]
        if not valid_confs:
            return 0
        return min(100, max(0, mean(valid_confs) * 0.9))

    @staticmethod
    def calculate_text_quality(text):
        if not text:
            return 0

        # Проверка на "мусорный" текст
        if len(text) > 1000 and len(set(text)) < 10:
            return 0

        metrics = {
            'unusual_chars': 0,
            'digit_ratio': 0,
            'line_breaks': 0,
            'word_lengths': [],
            'spaces_ratio': 0
        }

        words = text.split()
        metrics['word_lengths'] = [len(word) for word in words if word.isalpha()]
        metrics['spaces_ratio'] = text.count(' ') / len(text) if text else 0

        for char in text:
            if not char.isalnum() and char not in ' .,!?;:-()\n\'"':
                metrics['unusual_chars'] += 1
            elif char.isdigit():
                metrics['digit_ratio'] += 1

        metrics['digit_ratio'] = metrics['digit_ratio'] / len(text) if text else 0
        metrics['line_breaks'] = text.count('\n') / len(text.splitlines()) if text else 0

        quality = 100

        # Штрафы за проблемы
        quality -= metrics['unusual_chars'] * 0.5
        quality -= metrics['digit_ratio'] * 30
        quality -= min(20, metrics['line_breaks'] * 10)

        # Проверка среднего длины слова
        if metrics['word_lengths']:
            avg_word_len = mean(metrics['word_lengths'])
            if avg_word_len < 2 or avg_word_len > 12:
                quality -= 20

        # Проверка соотношения пробелов
        if not 0.1 < metrics['spaces_ratio'] < 0.3:
            quality -= 15

        return max(0, min(100, quality))


class OCRProcessor:
    def __init__(self):
        self.default_languages = "eng+rus"
        self.temp_dir = mkdtemp(prefix="ocr_temp_")
        self.accuracy_calculator = AccuracyCalculator()

    def __del__(self):
        try:
            for filename in os.listdir(self.temp_dir):
                file_path = os.path.join(self.temp_dir, filename)
                try:
                    if os.path.isfile(file_path):
                        os.unlink(file_path)
                except Exception as e:
                    logger.warning(f"Could not delete temp file {file_path}: {e}")
        except:
            pass

    def _is_image_data(self, data):
        """Проверяет, являются ли данные изображением"""
        try:
            Image.open(io.BytesIO(data))
            return True
        except UnidentifiedImageError:
            return False
        except Exception:
            return False

    def process_image(self, image):
        """Улучшенная обработка изображений с дополнительными этапами предобработки"""
        try:
            if isinstance(image, str):
                # Если передан путь к файлу
                img = cv2.imread(image)
                if img is None:
                    with Image.open(image) as img_pil:
                        img = cv2.cvtColor(np.array(img_pil), cv2.COLOR_RGB2BGR)
            else:
                # Если передано изображение напрямую
                img = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)

            if img is None:
                raise ValueError("Invalid image")

            # Конвертация в grayscale
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

            # Улучшение контраста с CLAHE
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            enhanced = clahe.apply(gray)

            # Шумоподавление
            denoised = cv2.fastNlMeansDenoising(enhanced, h=10)

            # Бинаризация
            _, binary = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

            # Дополнительная обработка для улучшения качества
            kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]])
            processed = cv2.filter2D(binary, -1, kernel)

            # Распознавание с улучшенными параметрами
            custom_config = r'--oem 3 --psm 6 -c preserve_interword_spaces=1'
            data = pytesseract.image_to_data(
                processed,
                lang=self.default_languages,
                config=custom_config,
                output_type=pytesseract.Output.DICT
            )

            text = ' '.join([word for word, conf in zip(data['text'], data['conf']) if int(conf) > 0])
            confidences = [float(c) for c in data['conf'] if float(c) > 0]

            accuracy = self.accuracy_calculator.calculate_from_confidence(confidences)
            text_quality = self.accuracy_calculator.calculate_text_quality(text)
            final_accuracy = (accuracy * 0.7 + text_quality * 0.3)

            return {
                'text': text.strip(),
                'accuracy': round(final_accuracy, 1),
                'confidence_data': confidences,
                'source': 'image'
            }
        except Exception as e:
            logger.error(f"OCR processing error: {str(e)}", exc_info=True)
            return {
                'text': f"Ошибка обработки: {str(e)}",
                'accuracy': 0,
                'confidence_data': [],
                'source': 'image'
            }

    def _extract_images_from_docx(self, docx_path):
        try:
            image_results = {
                'text': '',
                'accuracy': 0,
                'confidence_data': [],
                'image_count': 0
            }

            with zipfile.ZipFile(docx_path) as z:
                for file in z.namelist():
                    if file.startswith('word/media/') and file.split('.')[-1].lower() in ['png', 'jpg', 'jpeg', 'bmp']:
                        with z.open(file) as img_file:
                            img_data = img_file.read()

                            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_img:
                                temp_img.write(img_data)
                                temp_img_path = temp_img.name

                            try:
                                ocr_result = self.process_image(temp_img_path)
                                if ocr_result['text'].strip():
                                    image_results['text'] += f"{ocr_result['text']}\n"
                                    image_results['confidence_data'].extend(ocr_result['confidence_data'])
                                    image_results['image_count'] += 1
                            finally:
                                os.unlink(temp_img_path)

            if image_results['confidence_data']:
                image_results['accuracy'] = self.accuracy_calculator.calculate_from_confidence(
                    image_results['confidence_data']
                )

            return image_results
        except Exception as e:
            logger.error(f"Error extracting images from DOCX: {str(e)}", exc_info=True)
            return {
                'text': f"[Ошибка извлечения изображений: {str(e)}]",
                'accuracy': 0,
                'confidence_data': [],
                'image_count': 0
            }

    def _process_docx(self, docx_path):
        try:
            result = {
                'text': '',
                'accuracy': 0,
                'source': 'docx',
                'confidence_data': [],
                'has_text': False,
                'has_images': False
            }

            # Извлекаем текст из документа
            doc = Document(docx_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            doc_text = "\n".join(text_parts)
            result['has_text'] = bool(doc_text.strip())

            # Извлекаем текст из таблиц
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_texts.append(cell.text)

            if table_texts:
                doc_text += "\n\n[Текст из таблиц]:\n" + "\n".join(table_texts)
                result['has_text'] = True

            # Извлекаем текст из изображений
            image_result = self._extract_images_from_docx(docx_path)
            result['has_images'] = bool(image_result['text'].strip())

            # Комбинируем результаты
            result['text'] = doc_text
            if image_result['text'].strip():
                result['text'] += f"\n\n[Текст из изображений]:\n{image_result['text']}"
                result['confidence_data'] = image_result['confidence_data']

            # Рассчитываем точность
            if result['has_text'] and result['has_images']:
                # Если есть и текст и изображения
                text_quality = self.accuracy_calculator.calculate_text_quality(doc_text)
                combined_accuracy = (text_quality * 0.7 + image_result['accuracy'] * 0.3)
                result['accuracy'] = round(combined_accuracy, 1)
            elif result['has_text']:
                # Если только текст
                result['accuracy'] = self.accuracy_calculator.calculate_text_quality(doc_text)
            elif result['has_images']:
                # Если только изображения
                result['accuracy'] = image_result['accuracy']
            else:
                # Если ничего не найдено
                result['accuracy'] = 0
                result['text'] = "Не удалось извлечь текст из DOCX файла"

            return result
        except Exception as e:
            logger.error(f"DOCX processing error: {str(e)}", exc_info=True)
            return {
                'text': f"Ошибка обработки DOCX: {str(e)}",
                'accuracy': 0,
                'source': 'docx',
                'confidence_data': []
            }

    def _process_doc(self, doc_path):
        if not HAS_WIN32COM:
            return {
                'text': "Ошибка: Для обработки DOC файлов требуется установить pywin32\n"
                        "Установите: pip install pywin32",
                'accuracy': 0,
                'source': 'doc',
                'confidence_data': []
            }

        try:
            with com_lock:
                pythoncom.CoInitialize()

                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False

                try:
                    doc = word.Documents.Open(os.path.abspath(doc_path))
                    doc_text = doc.Content.Text.strip()

                    temp_pdf = os.path.join(self.temp_dir, f"temp_{os.path.basename(doc_path)}.pdf")
                    doc.SaveAs(temp_pdf, FileFormat=17)

                    pdf_result = self.process_pdf(temp_pdf)

                    result = {
                        'text': doc_text,
                        'accuracy': 95,
                        'source': 'doc',
                        'confidence_data': []
                    }

                    if pdf_result['text'].strip():
                        result['text'] += f"\n\n[Текст из изображений]:\n{pdf_result['text']}"
                        result['confidence_data'] = pdf_result.get('confidence_data', [])
                        result['accuracy'] = (95 * 0.7 + pdf_result['accuracy'] * 0.3)

                    return result

                finally:
                    doc.Close(False)
                    word.Quit()
                    if os.path.exists(temp_pdf):
                        os.remove(temp_pdf)

        except Exception as e:
            logger.error(f"DOC processing error: {str(e)}", exc_info=True)
            return {
                'text': f"Ошибка обработки DOC: {str(e)}",
                'accuracy': 0,
                'source': 'doc',
                'confidence_data': []
            }
        finally:
            pythoncom.CoUninitialize()

    def process_pdf(self, file_path):
        try:
            result = {
                'text': '',
                'accuracy': 0,
                'pages': [],
                'source': 'pdf',
                'confidence_data': []
            }

            try:
                with pdfplumber.open(file_path) as pdf:
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        if page_text:
                            page_quality = self.accuracy_calculator.calculate_text_quality(page_text)
                            result['pages'].append({
                                'text': page_text,
                                'accuracy': page_quality,
                                'method': 'direct_extract'
                            })
            except Exception as e:
                logger.warning(f"PDF text extraction failed, trying OCR: {e}")

            if not result['pages']:
                images = convert_from_path(
                    file_path,
                    dpi=300,
                    poppler_path=POPPLER_PATH,
                    output_folder=self.temp_dir,
                    fmt='jpeg'
                )

                for i, image in enumerate(images):
                    try:
                        temp_img_path = os.path.join(self.temp_dir, f"page_{i}.jpg")
                        image.save(temp_img_path, 'JPEG', quality=90)

                        page_result = self.process_image(temp_img_path)
                        result['pages'].append({
                            'text': page_result['text'],
                            'accuracy': page_result['accuracy'],
                            'method': 'ocr',
                            'confidence_data': page_result['confidence_data']
                        })

                        os.unlink(temp_img_path)
                    except Exception as e:
                        logger.error(f"Page {i} processing failed: {e}")
                        continue

            if result['pages']:
                result['text'] = "\n".join(
                    f"--- Страница {i + 1} [Точность: {p['accuracy']}%] ---\n{p['text']}"
                    for i, p in enumerate(result['pages'])
                )
                result['accuracy'] = round(mean(p['accuracy'] for p in result['pages']), 1)
                result['confidence_data'] = [c for p in result['pages'] for c in p.get('confidence_data', [])]

            return result
        except Exception as e:
            logger.error(f"PDF processing error: {str(e)}", exc_info=True)
            return {
                'text': f"Ошибка обработки PDF: {str(e)}",
                'accuracy': 0,
                'pages': [],
                'source': 'pdf',
                'confidence_data': []
            }

    def _process_txt(self, file_path):
        try:
            with open(file_path, 'rb') as f:
                content = f.read()

            # Проверяем, является ли файл изображением с расширением .txt
            if self._is_image_data(content):
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_img:
                    temp_img.write(content)
                    temp_path = temp_img.name

                img_result = self.process_image(temp_path)
                os.unlink(temp_path)

                return {
                    'text': img_result['text'],
                    'accuracy': img_result['accuracy'],
                    'source': 'image_in_txt',
                    'confidence_data': img_result.get('confidence_data', [])
                }

            # Если это обычный текстовый файл
            try:
                # Пробуем разные кодировки
                encodings = ['utf-8', 'windows-1251', 'cp866', 'iso-8859-5']
                text = None

                for encoding in encodings:
                    try:
                        text = content.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue

                if text is None:
                    return {
                        'text': "Не удалось определить кодировку текстового файла",
                        'accuracy': 0,
                        'source': 'text',
                        'confidence_data': []
                    }

                text = text.strip()
                quality = self.accuracy_calculator.calculate_text_quality(text)

                return {
                    'text': text,
                    'accuracy': quality,
                    'source': 'text',
                    'confidence_data': []
                }

            except Exception as e:
                logger.error(f"Text decoding error: {str(e)}", exc_info=True)
                return {
                    'text': f"Ошибка декодирования текста: {str(e)}",
                    'accuracy': 0,
                    'source': 'text',
                    'confidence_data': []
                }

        except Exception as e:
            logger.error(f"Text file processing error: {str(e)}", exc_info=True)
            return {
                'text': f"Ошибка обработки текстового файла: {str(e)}",
                'accuracy': 0,
                'source': 'text',
                'confidence_data': []
            }

    def process_file(self, file_path, file_extension):
        processors = {
            'pdf': self.process_pdf,
            'doc': self._process_doc,
            'docx': self._process_docx,
            'txt': self._process_txt,
            'jpg': self.process_image,
            'jpeg': self.process_image,
            'png': self.process_image,
            'tiff': self.process_image,
            'bmp': self.process_image,
            'webp': self.process_image
        }

        if file_extension not in processors:
            raise ValueError(f"Unsupported file type: {file_extension}")

        return processors[file_extension](file_path)


ocr_processor = OCRProcessor()


class FileProcessor:
    @staticmethod
    def save_to_db(file_data, session_id):
        file_data['session_id'] = session_id
        return db.files.insert_one(file_data).inserted_id

    @staticmethod
    def get_file_history(session_id, limit=20):
        files = list(db.files.find({'session_id': session_id}).sort('processed_date', -1).limit(limit))
        for file in files:
            file['_id'] = str(file['_id'])
        return files

    @staticmethod
    def get_file_by_id(file_id, session_id):
        file = db.files.find_one({'_id': ObjectId(file_id), 'session_id': session_id})
        if file:
            file['_id'] = str(file['_id'])
        return file

    @staticmethod
    def clear_history(session_id):
        return db.files.delete_many({'session_id': session_id})


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def get_file_size_string(size_bytes):
    if size_bytes < 1024:
        return f"{size_bytes} B"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    else:
        return f"{size_bytes / (1024 * 1024):.1f} MB"


@app.before_request
def assign_session():
    if 'session_id' not in session:
        session['session_id'] = os.urandom(16).hex()


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/history')
def history():
    files = FileProcessor.get_file_history(session['session_id'])
    for file in files:
        file['type'] = file['name'].split('.')[-1].upper()
        file['size'] = file['size_str']
        file['processed_date'] = file['processed_date'].strftime('%d %B %Y, %H:%M')
    return render_template('history.html', files=files)


@app.route('/file/<file_id>')
def file_details(file_id):
    file = FileProcessor.get_file_by_id(file_id, session['session_id'])
    if not file:
        flash('Файл не найден', 'error')
        return redirect(url_for('history'))

    file['type'] = file['name'].split('.')[-1].upper()
    file['size'] = file['size_str']
    file['processed_date'] = file['processed_date'].strftime('%d %B %Y, %H:%M')

    return render_template('file_details.html', file=file)


@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'Файл не выбран'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Файл не выбран'}), 400

    if file and allowed_file(file.filename):
        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)

        if file_size > app.config['MAX_CONTENT_LENGTH']:
            return jsonify({'error': 'Размер файла не должен превышать 10 MB'}), 400

        filename = secure_filename(file.filename)
        file_extension = filename.split('.')[-1].lower()
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        try:
            start_time = time.time()
            processing_result = ocr_processor.process_file(file_path, file_extension)
            processing_time = round(time.time() - start_time, 2)

            file_data = {
                'name': filename,
                'type': file_extension.upper(),
                'size': file_size,
                'size_str': get_file_size_string(file_size),
                'processed_date': datetime.now(),
                'processing_time': processing_time,
                'extracted_text': processing_result['text'],
                'accuracy': processing_result['accuracy'],
                'status': 'completed',
                'source_type': processing_result.get('source', 'unknown'),
                'summary': processing_result['text'][:100] + '...' if len(processing_result['text']) > 100 else
                processing_result['text']
            }

            if 'pages' in processing_result:
                file_data['page_count'] = len(processing_result['pages'])
                file_data['avg_page_accuracy'] = processing_result['accuracy']

            file_id = FileProcessor.save_to_db(file_data, session['session_id'])

            return jsonify({
                'success': True,
                'message': 'Файл обработан и добавлен в историю',
                'file_id': str(file_id),
                'filename': filename,
                'size': get_file_size_string(file_size),
                'processing_time': processing_time,
                'extracted_text': processing_result['text'],
                'accuracy': processing_result['accuracy'],
                'summary': file_data['summary']
            })

        except Exception as e:
            logger.error(f"Error processing file: {str(e)}", exc_info=True)
            return jsonify({'error': f'Ошибка обработки файла: {str(e)}'}), 500

        finally:
            if os.path.exists(file_path):
                os.remove(file_path)
    else:
        return jsonify({'error': 'Неподдерживаемый тип файла'}), 400


@app.route('/clear_history', methods=['POST'])
def clear_history():
    try:
        result = FileProcessor.clear_history(session['session_id'])
        flash(f'Удалено {result.deleted_count} файлов из истории', 'success')
    except Exception as e:
        logger.error(f"Error clearing history: {str(e)}", exc_info=True)
        flash('Ошибка при очистке истории', 'error')
    return redirect(url_for('history'))


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5001)